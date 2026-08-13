"""
Flask web application for OpenTale — powered by LangChain
"""

import json
import math
import os
import re
from io import BytesIO

from flask import (
    Flask,
    Response,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    stream_with_context,
)

import prompts
from agents import (
    PROMPT_DEBUGGING_DIR,
    REVISION_AGENTS,
    REVISION_AGENT_LABELS,
    BookAgents,
    check_openai_connection,
)
from config import get_config, save_config as save_config_file
from config import DEFAULT_CONFIG

# ============================================================
# Book / Project path management
# ============================================================

BASE_BOOK_DIR = "book_output"
TEXT_EXTENSION = ".txt"

# These globals are recomputed by reload_paths()
BOOK_OUTPUT_DIR = ""
WORLD_FILE = ""
CHARACTERS_FILE = ""
SYNOPSIS_FILE = ""
OUTLINE_FILE = ""
CHAPTERS_JSON_FILE = ""
MASTER_PROMPT_FILE = ""
SETTINGS_FILE = ""
OUTLINE_JSON_FILE = ""
CHAPTERS_DIR = ""


def reload_paths(project_name: str):
    """Recompute all path globals for the given project name."""
    global BOOK_OUTPUT_DIR, WORLD_FILE, CHARACTERS_FILE, SYNOPSIS_FILE
    global OUTLINE_FILE, CHAPTERS_JSON_FILE, MASTER_PROMPT_FILE, SETTINGS_FILE
    global OUTLINE_JSON_FILE, CHAPTERS_DIR

    BOOK_OUTPUT_DIR = os.path.join(BASE_BOOK_DIR, project_name)
    WORLD_FILE = os.path.join(BOOK_OUTPUT_DIR, f"world{TEXT_EXTENSION}")
    CHARACTERS_FILE = os.path.join(BOOK_OUTPUT_DIR, f"characters{TEXT_EXTENSION}")
    SYNOPSIS_FILE = os.path.join(BOOK_OUTPUT_DIR, f"synopsis{TEXT_EXTENSION}")
    OUTLINE_FILE = os.path.join(BOOK_OUTPUT_DIR, f"outline{TEXT_EXTENSION}")
    CHAPTERS_JSON_FILE = os.path.join(BOOK_OUTPUT_DIR, "chapters.json")
    MASTER_PROMPT_FILE = os.path.join(BOOK_OUTPUT_DIR, f"master_prompt{TEXT_EXTENSION}")
    SETTINGS_FILE = os.path.join(BOOK_OUTPUT_DIR, "settings.json")
    OUTLINE_JSON_FILE = os.path.join(BOOK_OUTPUT_DIR, "outline.json")
    CHAPTERS_DIR = os.path.join(BOOK_OUTPUT_DIR, "chapters")


def get_current_project_name() -> str:
    """Read the current project name from settings, or default."""
    project_file = os.path.join(BASE_BOOK_DIR, ".active_project")
    if os.path.exists(project_file):
        with open(project_file, "r", encoding="utf-8") as f:
            name = f.read().strip()
            if name:
                return name
    return "default"


def set_current_project_name(name: str):
    """Persist the active project name."""
    os.makedirs(BASE_BOOK_DIR, exist_ok=True)
    project_file = os.path.join(BASE_BOOK_DIR, ".active_project")
    with open(project_file, "w", encoding="utf-8") as f:
        f.write(name.strip())


def list_projects() -> list:
    """List all available projects (subdirectories of book_output/)."""
    if not os.path.exists(BASE_BOOK_DIR):
        return ["default"]
    entries = sorted(os.listdir(BASE_BOOK_DIR))
    projects = [e for e in entries if os.path.isdir(os.path.join(BASE_BOOK_DIR, e))]
    if "default" not in projects:
        projects.insert(0, "default")
    projects = [p for p in projects if not p.startswith(".") and p != "__pycache__"]
    return projects


# Initialise paths from the persisted project name
_current_project = get_current_project_name()
reload_paths(_current_project)

# Ensure directories exist
os.makedirs(CHAPTERS_DIR, exist_ok=True)

# ============================================================

PREVIOUS_CHAPTER_CONTEXT_LENGTH = 8000

app = Flask(__name__)
app.secret_key = "ai-book-writer-secret-key"


@app.context_processor
def inject_project():
    return {"current_project": _current_project, "projects": list_projects()}


# Initialize global variables
agent_config = get_config()


# ------------------------------------------------------------------
# Helper: format a LangChain ChatPromptTemplate and return the user message string
# ------------------------------------------------------------------
def _fmt_user(prompt_template, **kwargs) -> str:
    """Invoke a ChatPromptTemplate and return the first (user) message content as a string."""
    msgs = prompt_template.invoke(kwargs)
    return msgs.messages[0].content if msgs.messages else ""


# Helper functions to read data from files
def get_world_theme():
    """Get world theme from file."""
    if os.path.exists(WORLD_FILE):
        with open(WORLD_FILE, "r", encoding="utf-8") as f:
            return f.read().strip()
    return ""


def get_characters():
    """Get characters from file."""
    if os.path.exists(CHARACTERS_FILE):
        with open(CHARACTERS_FILE, "r", encoding="utf-8") as f:
            return f.read().strip()
    return ""


def get_outline():
    """Get outline from file."""
    if os.path.exists(OUTLINE_FILE):
        with open(OUTLINE_FILE, "r", encoding="utf-8") as f:
            return f.read().strip()
    return ""


def get_synopsis():
    """Get synopsis from file."""
    if os.path.exists(SYNOPSIS_FILE):
        with open(SYNOPSIS_FILE, "r", encoding="utf-8") as f:
            return f.read().strip()
    return ""


def get_chapters():
    """Get chapters from file, including a flag if content exists."""
    chapters = []
    if os.path.exists(CHAPTERS_JSON_FILE):
        with open(CHAPTERS_JSON_FILE, "r", encoding="utf-8") as f:
            try:
                chapters = json.load(f)
            except json.JSONDecodeError:
                chapters = []

    for chapter in chapters:
        chapter_file_path = os.path.join(
            CHAPTERS_DIR, f"chapter_{chapter['chapter_number']}{TEXT_EXTENSION}"
        )
        chapter["has_content"] = (
            os.path.exists(chapter_file_path) and os.path.getsize(chapter_file_path) > 0
        )
        editor_chapter_file_path = os.path.join(
            CHAPTERS_DIR, f"chapter_{chapter['chapter_number']}_editor{TEXT_EXTENSION}"
        )
        chapter["has_been_reviewed"] = (
            os.path.exists(editor_chapter_file_path)
            and os.path.getsize(editor_chapter_file_path) > 0
        )
        action_beats_file_path = os.path.join(
            CHAPTERS_DIR,
            f"chapter_{chapter['chapter_number']}_action_beats{TEXT_EXTENSION}",
        )
        chapter["has_action_beats"] = (
            os.path.exists(action_beats_file_path)
            and os.path.getsize(action_beats_file_path) > 0
        )
    return chapters


def get_paginated_chapters(page, per_page):
    """Helper to get paginated chapters."""
    all_chapters = get_chapters()
    total_chapters = len(all_chapters)
    total_pages = math.ceil(total_chapters / per_page) if per_page > 0 else 1

    start_index = (page - 1) * per_page
    end_index = start_index + per_page
    paginated_chapters = all_chapters[start_index:end_index]

    return {
        "chapters": paginated_chapters,
        "total_pages": total_pages,
        "current_page": page,
        "total_chapters": total_chapters,
        "per_page": per_page,
    }


def get_paginated_chapters_from_request(request, chapters, chapter_number):
    chapters_per_page = request.args.get("per_page", 10, type=int)

    if "page" in request.args:
        page = request.args.get("page", 1, type=int)
    else:
        try:
            active_chapter_index = [c["chapter_number"] for c in chapters].index(
                chapter_number
            )
            page = math.ceil((active_chapter_index + 1) / chapters_per_page)
        except (ValueError, ZeroDivisionError):
            page = 1

    chapters_paginated = get_paginated_chapters(page, chapters_per_page)
    return chapters_paginated


def get_action_beats(chapter_number):
    """Get action beats for a specific chapter from file."""
    action_beats_path = os.path.join(
        CHAPTERS_DIR, f"chapter_{chapter_number}_action_beats{TEXT_EXTENSION}"
    )
    if os.path.exists(action_beats_path):
        with open(action_beats_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    return ""


def get_master_prompt():
    """Get master prompt from file."""
    if os.path.exists(MASTER_PROMPT_FILE):
        with open(MASTER_PROMPT_FILE, "r", encoding="utf-8") as f:
            return f.read().strip()
    return ""


def get_settings():
    """Get settings from file."""
    if os.path.exists(SETTINGS_FILE):
        with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return {}
    return {}


def get_previous_chapter_context(chapter_number):
    """Get context from the previous chapter to ensure continuity."""
    previous_context = ""
    if chapter_number > 1:
        prev_editor_path = os.path.join(
            CHAPTERS_DIR,
            f"chapter_{chapter_number - 1}_editor{TEXT_EXTENSION}",
        )
        prev_chapter_path = os.path.join(
            CHAPTERS_DIR, f"chapter_{chapter_number - 1}{TEXT_EXTENSION}"
        )
        prev_path = None
        if os.path.exists(prev_editor_path) and os.path.getsize(prev_editor_path) > 0:
            prev_path = prev_editor_path
        elif os.path.exists(prev_chapter_path):
            prev_path = prev_chapter_path

        if prev_path:
            with open(prev_path, "r", encoding="utf-8") as f:
                content = f.read()
                previous_context = content[-PREVIOUS_CHAPTER_CONTEXT_LENGTH:]

    return previous_context


def save_settings(settings):
    """Save settings to file."""
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(settings, f, indent=2)


# ==================================================================
# Routes
# ==================================================================


@app.route("/")
def index():
    """Render the home page"""
    chapters = get_chapters()
    return render_template("index.html", chapters=chapters)


@app.route("/config", methods=["GET"])
def config():
    """Display config interface"""
    from config import _load_config_file
    settings = get_config()
    raw_settings = _load_config_file()
    return render_template("config.html", settings=settings, raw_settings=raw_settings)


@app.route("/save_config", methods=["POST"])
def save_config():
    """Save configuration from the editable config form."""
    from config import save_config as save_config_file
    data = request.json
    if not data:
        return jsonify({"success": False, "error": "No data provided"}), 400
    save_config_file(data)
    global agent_config
    agent_config = get_config()
    return jsonify({"success": True, "message": "Configuration saved."})


@app.route("/check_connection", methods=["POST"])
def check_connection():
    """Test the AI API connection."""
    data = request.json or {}
    base_url = data.get("base_url") or agent_config["config_list"][0]["base_url"]
    api_key = data.get("api_key") or agent_config["config_list"][0]["api_key"]
    try:
        from openai import OpenAI
        client = OpenAI(base_url=base_url, api_key=api_key)
        client.models.list()
        return jsonify({"success": True, "message": "Connection successful! API is reachable."})
    except Exception as e:
        return jsonify({"success": False, "message": f"Connection failed: {str(e)}"})


# ---- Project switching endpoints ----

@app.route("/set_project", methods=["POST"])
def set_project():
    """Switch to a different project (book). Reloads all paths."""
    global _current_project
    data = request.json
    name = data.get("project", "default").strip()
    if not name or name.startswith("."):
        return jsonify({"success": False, "error": "Invalid project name"}), 400
    _current_project = name
    set_current_project_name(name)
    reload_paths(name)
    os.makedirs(CHAPTERS_DIR, exist_ok=True)
    return jsonify({"success": True, "project": name})


@app.route("/create_project", methods=["POST"])
def create_project():
    """Create a new project (empty book) and switch to it."""
    global _current_project
    data = request.json
    name = data.get("project", "").strip()
    if not name or name.startswith("."):
        return jsonify({"success": False, "error": "Invalid project name"}), 400
    project_dir = os.path.join(BASE_BOOK_DIR, name)
    if os.path.exists(project_dir):
        return jsonify({"success": False, "error": f"Project '{name}' already exists"}), 400
    os.makedirs(os.path.join(project_dir, "chapters"), exist_ok=True)
    _current_project = name
    set_current_project_name(name)
    reload_paths(name)
    return jsonify({"success": True, "project": name})


# ---- End project switching ----


@app.route("/synopsis", methods=["GET"])
def synopsis():
    if not os.path.exists("config.json"):
        flash("You need to create a config first.", "warning")
        return redirect("/config")
    synopsis_content = get_synopsis()
    settings = get_settings()
    chapters = get_chapters()
    return render_template(
        "synopsis.html",
        synopsis=synopsis_content,
        topic=settings.get("topic", ""),
        chapters=chapters,
    )


@app.route("/synopsis_chat_stream", methods=["POST"])
def synopsis_chat_stream():
    """Handle ongoing chat for synopsis building with streaming response"""
    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])
    topic = data.get("topic", "")

    if topic:
        settings = get_settings()
        settings["topic"] = topic
        save_settings(settings)

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(topic, 0)

    stream = book_agents.generate_chat_response_synopsis_stream(
        chat_history, topic, user_message
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/finalize_synopsis_stream", methods=["POST"])
def finalize_synopsis_stream():
    """Finalize the synopsis based on chat history with streaming response"""
    data = request.json
    chat_history = data.get("chat_history", [])
    topic = data.get("topic", "")

    if not chat_history:
        return jsonify({
            "error": "Chat history is empty. Please chat with the AI first to build your synopsis."
        }), 400

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(topic, 0)

    stream = book_agents.generate_final_synopsis_stream(chat_history, topic)

    def generate():
        yield 'data: {"content": ""}\n\n'
        collected_content = []
        for chunk in stream:
            content = chunk.content
            if content:
                collected_content.append(content)
                yield f"data: {json.dumps({'content': content})}\n\n"

        complete_content = "".join(collected_content)
        synopsis_content = complete_content.strip()
        synopsis_content = re.sub(r"\n+", "\n", synopsis_content)

        with open(SYNOPSIS_FILE, "w", encoding='utf-8') as f:
            f.write(synopsis_content)

        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/save_synopsis", methods=["POST"])
def save_synopsis():
    """Save edited synopsis"""
    synopsis_content = request.form.get("synopsis")
    with open(SYNOPSIS_FILE, "w", encoding='utf-8') as f:
        f.write(synopsis_content)
    return jsonify({"success": True})


@app.route("/world", methods=["GET"])
def world():
    if not os.path.exists(SYNOPSIS_FILE):
        flash("You need to create a synopsis first.", "warning")
        return redirect("/synopsis")

    world_theme = get_world_theme()
    settings = get_settings()
    chapters = get_chapters()
    return render_template(
        "world.html",
        world_theme=world_theme,
        topic=settings.get("topic", ""),
        chapters=chapters,
    )


@app.route("/world_chat", methods=["POST"])
def world_chat():
    """Handle ongoing chat for world building"""
    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])
    topic = data.get("topic", "")

    if topic:
        settings = get_settings()
        settings["topic"] = topic
        save_settings(settings)

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(topic, 0)

    ai_response = book_agents.generate_chat_response_world(
        chat_history, topic, user_message
    )
    ai_response = ai_response.strip()
    return jsonify({"message": ai_response})


@app.route("/world_chat_stream", methods=["POST"])
def world_chat_stream():
    """Handle ongoing chat for world building with streaming response"""
    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])
    topic = data.get("topic", "")

    if topic:
        settings = get_settings()
        settings["topic"] = topic
        save_settings(settings)

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(topic, 0)

    stream = book_agents.generate_chat_response_world_stream(
        chat_history, topic, user_message
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/finalize_world", methods=["POST"])
def finalize_world():
    """Finalize the world setting based on chat history"""
    data = request.json
    chat_history = data.get("chat_history", [])
    topic = data.get("topic", "")

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(topic, 0)

    world_theme = book_agents.generate_final_world(chat_history, topic)
    world_theme = world_theme.strip()
    world_theme = re.sub(r"\n+", "\n", world_theme.strip())

    with open(WORLD_FILE, "w", encoding='utf-8') as f:
        f.write(world_theme)

    return jsonify({"world_theme": world_theme})


@app.route("/finalize_world_stream", methods=["POST"])
def finalize_world_stream():
    """Finalize the world setting based on chat history with streaming response"""
    data = request.json
    chat_history = data.get("chat_history", [])
    topic = data.get("topic", "")

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(topic, 0)

    stream = book_agents.generate_final_world_stream(chat_history, topic)

    def generate():
        yield 'data: {"content": ""}\n\n'
        collected_content = []
        for chunk in stream:
            content = chunk.content
            if content:
                collected_content.append(content)
                yield f"data: {json.dumps({'content': content})}\n\n"

        complete_content = "".join(collected_content)
        world_theme = complete_content.strip()
        world_theme = re.sub(r"\n+", "\n", world_theme)

        with open(WORLD_FILE, "w", encoding='utf-8') as f:
            f.write(world_theme)

        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/save_world", methods=["POST"])
def save_world():
    """Save edited world theme"""
    world_theme = request.form.get("world_theme")
    with open(WORLD_FILE, "w", encoding='utf-8') as f:
        f.write(world_theme)
    return jsonify({"success": True})


@app.route("/characters", methods=["GET"])
def characters():
    if not os.path.exists(SYNOPSIS_FILE):
        flash("You need to create a synopsis first.", "warning")
        return redirect("/synopsis")

    characters_content = get_characters()
    world_theme = get_world_theme()
    synopsis = get_synopsis()
    chapters = get_chapters()
    settings = get_settings()
    num_characters = settings.get("num_characters", 3)

    return render_template(
        "characters.html",
        characters=characters_content,
        world_theme=world_theme,
        synopsis=synopsis,
        num_characters=num_characters,
        chapters=chapters,
    )


@app.route("/save_characters", methods=["POST"])
def save_characters():
    """Save edited characters"""
    characters_content = request.form.get("characters")
    with open(CHARACTERS_FILE, "w", encoding='utf-8') as f:
        f.write(characters_content)
    return jsonify({"success": True})


@app.route("/outline", methods=["GET", "POST"])
def outline():
    if not os.path.exists(SYNOPSIS_FILE):
        flash("You need to create a synopsis first.", "warning")
        return redirect("/synopsis")
    if not os.path.exists(WORLD_FILE):
        flash("You need to create a world setting first.", "warning")
        return redirect("/world")
    if not os.path.exists(CHARACTERS_FILE):
        flash("You need to create characters first.", "warning")
        return redirect("/characters")

    with open(WORLD_FILE, "r", encoding="utf-8") as f:
        world_theme = f.read()
    with open(CHARACTERS_FILE, "r", encoding="utf-8") as f:
        characters = f.read()
    with open(SYNOPSIS_FILE, "r", encoding="utf-8") as f:
        synopsis = f.read()

    outline_content = ""
    if os.path.exists(OUTLINE_FILE):
        with open(OUTLINE_FILE, "r", encoding="utf-8") as f:
            outline_content = f.read()

    chapters = get_chapters()
    settings = get_settings()
    num_chapters = settings.get("num_chapters", 20)

    return render_template(
        "outline.html",
        world_theme=world_theme,
        characters=characters,
        synopsis=synopsis,
        outline=outline_content,
        chapters=chapters,
        num_chapters=num_chapters,
    )


def _build_book_text():
    """Compile all chapters into a single book text (used for stats/diagnostics)."""
    collected = _collect_chapters()
    parts = [
        f"Chapter {ch['number']}: {ch['title']}\n\n{ch['content']}"
        for ch in collected
    ]
    included = [
        {"number": ch["number"], "title": ch["title"], "reviewed": ch["reviewed"]}
        for ch in collected
    ]
    return "\n\n\n".join(parts), included


def _build_book_docx():
    """Build a well-formatted .docx book from all chapters and return its bytes."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    collected = _collect_chapters()
    doc = Document()

    # Base body font (book-like serif)
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    # Title page
    title_text = get_current_project_name().replace("_", " ").replace("-", " ").title()
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f"A novel in {len(collected)} chapters"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = doc.styles["Subtitle"]

    doc.add_page_break()

    for index, chapter in enumerate(collected):
        heading = doc.add_heading(
            f"Chapter {chapter['number']}: {chapter['title']}", level=1
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Split content into logical paragraphs on blank lines.
        blocks = re.split(r"\n\s*\n", chapter["content"])
        for block in blocks:
            block = re.sub(r"\s*\n\s*", " ", block).strip()
            if block:
                doc.add_paragraph(block)

        # Page break between chapters (not after the last one).
        if index < len(collected) - 1:
            doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@app.route("/chapters", methods=["GET"])
def chapters_list():
    if not os.path.exists(SYNOPSIS_FILE):
        flash("You need to create a synopsis first.", "warning")
        return redirect("/synopsis")
    chapters = get_chapters()
    _, included = _build_book_text()
    included_numbers = {ch["number"] for ch in included}
    book_stats = {
        "total_chapters": len(chapters),
        "completed_chapters": len(included_numbers),
        "ready": len(included_numbers) > 0,
    }
    return render_template(
        "chapters.html", chapters=chapters, book_stats=book_stats
    )


@app.route("/export_book", methods=["GET"])
def export_book():
    """Generate the complete book from all chapters and serve it as a .docx download."""
    collected = _collect_chapters()
    if not collected:
        return jsonify({
            "error": "No chapter content found. Generate or review chapters first."
        }), 400

    project_name = get_current_project_name()
    filename = f"{project_name}.docx"

    try:
        docx_bytes = _build_book_docx()
    except Exception as e:
        return jsonify({
            "error": f"Failed to generate the Word document: {str(e)}"
        }), 500

    data = BytesIO(docx_bytes)
    return send_file(
        data,
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        as_attachment=True,
        download_name=filename,
    )


@app.route("/continuity", methods=["GET"])
def continuity():
    """Display the continuity check page."""
    if not os.path.exists(SYNOPSIS_FILE):
        flash("You need to create a synopsis first.", "warning")
        return redirect("/synopsis")
    chapters = get_chapters()
    return render_template("continuity.html", chapters=chapters)


def _clean_chapter_text(content: str) -> str:
    """Strip internal drafting markers (SCENE:, SCENE FINAL:, END OF CHAPTER,
    EDITED_SCENE:, etc.) from chapter text so only the prose remains.

    Handles markdown-wrapped markers (e.g. **SCENE:**), escaped underscores
    (EDITED\\_SCENE:), word-count tails (SCENE FINAL: 5078 words), and scene
    headings without destroying real prose.
    """
    content = content.replace("\r\n", "\n").replace("\r", "\n")

    out = []
    for raw in content.split("\n"):
        line = raw.strip()

        # Preserve blank lines: they separate paragraphs.
        if not line:
            out.append("")
            continue

        # Strip markdown emphasis characters and unescape underscores so we can
        # classify the line's actual content.
        core = line.strip(" *")
        core = core.replace("\\_", "_").strip()
        norm = re.sub(r"\s+", " ", core).strip()

        # --- Drop whole lines that are pure markers ---------------------------------
        if re.fullmatch(
            r"(SCENE\s*FINAL?\s*:?|EDITED\s*_?\s*SCENE\s*:?|END\s+OF\s+CHAPTER\s*\d*)",
            norm,
            re.IGNORECASE,
        ):
            continue

        # "SCENE FINAL: END OF CHAPTER N" / "END OF CHAPTER N"
        if re.fullmatch(
            r"SCENE\s*FINAL\s*:\s*END\s+OF\s+CHAPTER\s*\d*",
            norm,
            re.IGNORECASE,
        ):
            continue

        # "SCENE FINAL: 5078 words"
        if re.fullmatch(
            r"SCENE\s*FINAL\s*:\s*\d+\s+WORDS?",
            norm,
            re.IGNORECASE,
        ):
            continue

        # "SCENE FINAL: CHAPTER 2 - / – THE FIRST EXPERIMENT" (heading only)
        if re.fullmatch(
            r"SCENE\s*FINAL?\s*:\s*CHAPTER\s+\d+\s*[-\u2013].*",
            norm,
            re.IGNORECASE,
        ):
            continue

        # "SCENE: <text>" — strip the prefix; drop if it looks like a short
        # title/heading rather than prose.
        scene_match = re.fullmatch(r"SCENE\s*:\s*(.*)", norm, re.IGNORECASE)
        if scene_match:
            rest = scene_match.group(1).strip()
            if not rest:
                continue
            words = rest.split()
            # Short, non-sentence text (e.g. "The Glow of Discovery",
            # "Chapter 9", "THE SENSORY HYPOTHESIS") is a heading -> drop.
            if len(words) <= 6 and not re.search(r"[.!?]$", rest):
                continue
            line = rest
            out.append(line)
            continue

        # --- Remove inline "SCENE FINAL: ..." tails -------------------------------
        line = re.sub(
            r"\s*SCENE\s*FINAL\s*:\s*(END\s+OF\s+CHAPTER\s*\d*|\d+\s+WORDS?)",
            "",
            line,
            flags=re.IGNORECASE,
        )

        # Strip a leading "SCENE:" prefix when genuine prose follows.
        prefix_match = re.match(r"^SCENE\s*:\s+(.+)$", line, re.IGNORECASE)
        if prefix_match:
            line = prefix_match.group(1)

        if line.strip():
            out.append(line.strip())

    text = "\n".join(out)
    # Collapse runs of blank lines into a single blank line and trim the edges.
    text = re.sub(r"\n{3,}", "\n\n", text).strip("\n")
    return text.strip()


def _collect_chapters():
    """Collect all chapters with content, in order.

    For each chapter, prefer the reviewed ("_editor") version when available;
    otherwise fall back to the draft chapter. Returns a list of dicts with
    keys: number, title, content, reviewed.
    """
    chapters = get_chapters()
    chapters = sorted(chapters, key=lambda c: c.get("chapter_number", 0))

    collected = []
    for chapter in chapters:
        number = chapter.get("chapter_number")
        title = chapter.get("title", "")

        editor_path = os.path.join(CHAPTERS_DIR, f"chapter_{number}_editor{TEXT_EXTENSION}")
        draft_path = os.path.join(CHAPTERS_DIR, f"chapter_{number}{TEXT_EXTENSION}")

        content = ""
        used_editor = False
        if os.path.exists(editor_path) and os.path.getsize(editor_path) > 0:
            with open(editor_path, "r", encoding="utf-8") as f:
                content = f.read().strip()
            used_editor = True
        elif os.path.exists(draft_path) and os.path.getsize(draft_path) > 0:
            with open(draft_path, "r", encoding="utf-8") as f:
                content = f.read().strip()

        if not content:
            continue

        collected.append({
            "number": number,
            "title": title,
            "content": _clean_chapter_text(content),
            "reviewed": used_editor,
        })

    return collected


@app.route("/generate_chapters", methods=["POST"])
def generate_chapters():
    """Generate chapters structure from existing outline"""
    if not os.path.exists(OUTLINE_FILE):
        return jsonify({"error": "Outline not found. Please create an outline first."})

    with open(OUTLINE_FILE, "r", encoding="utf-8") as f:
        outline_content = f.read()

    num_chapters = int(request.form.get("num_chapters", 10))
    chapters = parse_outline_to_chapters(outline_content, num_chapters)

    with open(CHAPTERS_JSON_FILE, "w", encoding="utf-8") as f:
        json.dump(chapters, f, indent=2)

    return jsonify({"success": True, "num_chapters": len(chapters)})


@app.route("/save_outline", methods=["POST"])
def save_outline():
    """Save edited outline and generate chapters structure"""
    outline_content = request.form.get("outline")
    with open(OUTLINE_FILE, "w", encoding='utf-8') as f:
        f.write(outline_content)

    num_chapters = int(request.form.get("num_chapters", 10))
    chapters = parse_outline_to_chapters(outline_content, num_chapters)

    with open(CHAPTERS_JSON_FILE, "w", encoding="utf-8") as f:
        json.dump(chapters, f, indent=2)

    return jsonify({"success": True, "num_chapters": len(chapters)})


@app.route("/chapter/<int:chapter_number>", methods=["GET", "POST"])
def chapter(chapter_number):
    """Generate or display a specific chapter"""
    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )

    if not chapter_data:
        return render_template(
            "error.html", message=f"Chapter {chapter_number} not found"
        )

    if request.method == "POST":
        additional_context = request.form.get("additional_context", "")
        master_prompt = request.form.get("master_prompt", "")
        point_of_view = request.form.get("point_of_view", "Third-person limited")
        tense = request.form.get("tense", "Past tense")
        min_words = request.form.get("min_words", "5000")
        min_tokens = int(int(min_words) / 0.75)
        action_beats = request.form.get("action_beats_content", "")

        settings_to_save = get_settings()
        if "chapters" not in settings_to_save:
            settings_to_save["chapters"] = {}
        if str(chapter_number) not in settings_to_save["chapters"]:
            settings_to_save["chapters"][str(chapter_number)] = {}
        settings_to_save["chapters"][str(chapter_number)]["point_of_view"] = point_of_view
        settings_to_save["chapters"][str(chapter_number)]["tense"] = tense
        settings_to_save["chapters"][str(chapter_number)]["min_words"] = min_words
        save_settings(settings_to_save)

        world_theme = get_world_theme()
        characters = get_characters()
        previous_context = get_previous_chapter_context(chapter_number)

        book_agents = BookAgents(agent_config, chapters)
        book_agents.create_agents(world_theme, len(chapters))

        chapter_prompt = (
            f"{chapter_data['prompt']}\n\n{additional_context}"
            if additional_context
            else chapter_data["prompt"]
        )

        user_prompt_str = _fmt_user(
            prompts.CHAPTER_GENERATION_PROMPT,
            master_prompt=master_prompt,
            chapter_number=chapter_number,
            chapter_title=chapter_data["title"],
            chapter_outline=chapter_prompt,
            world_theme=world_theme,
            relevant_characters=characters,
            scene_details="",
            action_beats=action_beats,
            previous_context=previous_context,
            research_brief="",
            point_of_view=point_of_view,
            tense=tense,
            min_words=min_words,
            min_tokens=min_tokens,
        )

        chapter_content = book_agents.generate_content("writer", user_prompt_str)
        chapter_content = chapter_content.strip()
        chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{chapter_number}{TEXT_EXTENSION}")
        with open(chapter_path, "w", encoding="utf-8") as f:
            f.write(chapter_content)

        return jsonify({"chapter_content": chapter_content})

    # GET
    chapter_content = ""
    chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{chapter_number}{TEXT_EXTENSION}")
    if os.path.exists(chapter_path):
        with open(chapter_path, "r", encoding="utf-8") as f:
            chapter_content = f.read().strip()

    master_prompt = get_master_prompt()
    action_beats_content = get_action_beats(chapter_number)
    settings = get_settings()

    chapter_settings = settings.get("chapters", {}).get(str(chapter_number), {})
    point_of_view = chapter_settings.get("point_of_view", "Third-person limited")
    tense = chapter_settings.get("tense", "Past tense")
    min_words = chapter_settings.get("min_words", "5000")

    chapters_paginated = get_paginated_chapters_from_request(request, chapters, chapter_number)

    return render_template(
        "chapter.html",
        chapter=chapter_data,
        chapter_content=chapter_content,
        action_beats_content=action_beats_content,
        chapters=chapters,
        chapters_paginated=chapters_paginated,
        master_prompt=master_prompt,
        point_of_view=point_of_view,
        tense=tense,
        min_words=min_words,
    )


def _handle_chapter_stream(chapter_number, agent_name):
    """A helper function to handle chapter stream generation for both writer and editor."""

    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )

    if not chapter_data:
        return Response(
            json.dumps({"error": f"Chapter {chapter_number} not found"}),
            status=404,
            mimetype="application/json",
        )

    data = request.json
    additional_context = data.get("additional_context", "")
    master_prompt = data.get("master_prompt", "")
    point_of_view = data.get("point_of_view", "Third-person limited")
    tense = data.get("tense", "Past tense")
    min_words = data.get("min_words", "5000")
    min_tokens = int(int(min_words) / 0.75)
    action_beats = data.get("action_beats_content", "")
    show_prompt = data.get("show_prompt", False)
    chapter_content = data.get("chapter_content", "")

    settings_to_save = get_settings()
    if "chapters" not in settings_to_save:
        settings_to_save["chapters"] = {}
    if str(chapter_number) not in settings_to_save["chapters"]:
        settings_to_save["chapters"][str(chapter_number)] = {}
    settings_to_save["chapters"][str(chapter_number)]["point_of_view"] = point_of_view
    settings_to_save["chapters"][str(chapter_number)]["tense"] = tense
    settings_to_save["chapters"][str(chapter_number)]["min_words"] = min_words
    save_settings(settings_to_save)

    world_theme = get_world_theme()
    characters = get_characters()
    previous_context = get_previous_chapter_context(chapter_number)

    book_agents = BookAgents(agent_config, chapters)
    book_agents.create_agents(world_theme, len(chapters))

    chapter_prompt = (
        f"{chapter_data['prompt']}\n\n{additional_context}"
        if additional_context
        else chapter_data["prompt"]
    )

    # Auto-fetch research brief if the research agent is enabled
    research_brief = ""
    if not show_prompt:
        from config import _load_config_file
        raw_cfg = _load_config_file()
        if raw_cfg.get("research_agent_enabled", False):
            searxng_host = raw_cfg.get("searxng_host", "").strip()
            if searxng_host:
                try:
                    search_results = book_agents.fetch_research_results(
                        searxng_host,
                        chapter_data["title"],
                        chapter_prompt,
                        world_theme,
                        characters,
                    )
                    brief = book_agents.generate_research_brief(
                        chapter_data["title"],
                        chapter_prompt,
                        world_theme,
                        characters,
                        search_results,
                    )
                    research_brief = brief
                    # Save per-chapter brief
                    brief_path = os.path.join(
                        CHAPTERS_DIR,
                        f"chapter_{chapter_number}_research_brief{TEXT_EXTENSION}",
                    )
                    with open(brief_path, "w", encoding="utf-8") as f:
                        f.write(brief)
                except Exception:
                    # Research should never block chapter generation
                    research_brief = ""

    # ------------------------------------------------------------------
    # Editor -> multi-agent revision pipeline (the 8 specialized agents)
    # ------------------------------------------------------------------
    if agent_name == "editor":
        revision_context = {
            "chapter_number": chapter_number,
            "chapter_title": chapter_data["title"],
            "chapter_content": chapter_content,
            "chapter_outline": chapter_prompt,
            "world_theme": world_theme,
            "characters": characters,
            "action_beats": action_beats,
            "previous_context": previous_context,
            "research_brief": research_brief,
            "master_prompt": master_prompt,
            "point_of_view": point_of_view,
            "tense": tense,
            "min_words": min_words,
        }

        if show_prompt:
            pipeline_desc = "Chapter Review Pipeline (multi-agent):\n\n" + "\n".join(
                f"{i}. {REVISION_AGENT_LABELS.get(name, name)}"
                for i, name in enumerate(REVISION_AGENTS, 1)
            )
            first_agent = REVISION_AGENTS[0]
            system_prompt = (
                pipeline_desc
                + "\n\n--- First agent system prompt ---\n\n"
                + book_agents.system_prompts.get(first_agent, "")
            )
            user_prompt = book_agents._build_revision_prompt(
                chapter_number=chapter_number,
                chapter_title=chapter_data["title"],
                chapter_content=chapter_content,
                chapter_outline=chapter_prompt,
                world_theme=world_theme,
                characters=characters,
                action_beats=action_beats,
                previous_context=previous_context,
                research_brief=research_brief,
                master_prompt=master_prompt,
                point_of_view=point_of_view,
                tense=tense,
                min_words=min_words,
            )
            return jsonify({"system_prompt": system_prompt, "user_prompt": user_prompt})

        def generate():
            yield 'data: {"content": ""}\n\n'
            collected_content = []
            for payload in book_agents.revise_chapter_stream(revision_context):
                if "content" in payload:
                    collected_content.append(payload["content"])
                yield f"data: {json.dumps(payload)}\n\n"

            complete_content = "".join(collected_content)
            chapter_path = os.path.join(
                CHAPTERS_DIR, f"chapter_{chapter_number}_editor{TEXT_EXTENSION}"
            )
            with open(chapter_path, "w", encoding="utf-8") as f:
                f.write(complete_content)

            yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

        return Response(
            stream_with_context(generate()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    # ------------------------------------------------------------------
    # Writer -> single-agent chapter generation (unchanged)
    # ------------------------------------------------------------------
    user_prompt_str = _fmt_user(
        prompts.CHAPTER_GENERATION_PROMPT,
        master_prompt=master_prompt,
        chapter_number=chapter_number,
        chapter_title=chapter_data["title"],
        chapter_outline=chapter_prompt,
        world_theme=world_theme,
        relevant_characters=characters,
        scene_details="",
        action_beats=action_beats,
        previous_context=previous_context,
        research_brief=research_brief,
        point_of_view=point_of_view,
        tense=tense,
        min_words=min_words,
        min_tokens=min_tokens,
    )

    if show_prompt:
        system_prompt = book_agents.system_prompts.get(agent_name, "")
        return jsonify({"system_prompt": system_prompt, "user_prompt": user_prompt_str})

    stream = book_agents.generate_content_stream(agent_name, user_prompt_str)

    def generate():
        yield 'data: {"content": ""}\n\n'
        collected_content = []
        for chunk in stream:
            content = chunk.content
            if content:
                collected_content.append(content)
                yield f"data: {json.dumps({'content': content})}\n\n"

        complete_content = "".join(collected_content)
        chapter_path = os.path.join(
            CHAPTERS_DIR, f"chapter_{chapter_number}{TEXT_EXTENSION}"
        )
        with open(chapter_path, "w", encoding="utf-8") as f:
            f.write(complete_content)

        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/chapter_stream/<int:chapter_number>", methods=["POST"])
def chapter_stream(chapter_number):
    """Generate or display a specific chapter using the writer agent."""
    return _handle_chapter_stream(chapter_number, "writer")


@app.route("/chapter_editor/<int:chapter_number>", methods=["GET"])
def chapter_editor(chapter_number):
    """Generate or display a specific chapter for editing"""
    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )

    if not chapter_data:
        return render_template(
            "error.html", message=f"Chapter {chapter_number} not found"
        )

    chapter_file_path = os.path.join(
        CHAPTERS_DIR, f"chapter_{chapter_number}{TEXT_EXTENSION}"
    )
    original_chapter_content = ""
    if os.path.exists(chapter_file_path):
        with open(chapter_file_path, "r", encoding="utf-8") as f:
            original_chapter_content = f.read()

    editor_review_file_path = os.path.join(
        CHAPTERS_DIR, f"chapter_{chapter_number}_editor{TEXT_EXTENSION}"
    )
    chapter_content = ""
    has_review = False
    if os.path.exists(editor_review_file_path):
        with open(editor_review_file_path, "r", encoding="utf-8") as f:
            chapter_content = f.read()
        has_review = True

    previous_context = get_previous_chapter_context(chapter_number)
    master_prompt = get_master_prompt()
    action_beats_content = get_action_beats(chapter_number)
    settings = get_settings()

    chapter_settings = settings.get("chapters", {}).get(str(chapter_number), {})
    point_of_view = chapter_settings.get("point_of_view", "Third-person limited")
    tense = chapter_settings.get("tense", "Past tense")
    min_words = chapter_settings.get("min_words", "5000")
    min_tokens = int(int(min_words) / 0.75)

    chapters_paginated = get_paginated_chapters_from_request(request, chapters, chapter_number)

    return render_template(
        "chapter_editor.html",
        chapter=chapter_data,
        chapters=chapters,
        chapters_paginated=chapters_paginated,
        original_chapter_content=original_chapter_content,
        chapter_content=chapter_content,
        has_review=has_review,
        previous_context=previous_context,
        master_prompt=master_prompt,
        point_of_view=point_of_view,
        tense=tense,
        action_beats_content=action_beats_content,
        min_words=min_words,
        min_tokens=min_tokens,
    )


@app.route("/chapter_editor_stream/<int:chapter_number>", methods=["POST"])
def chapter_editor_stream(chapter_number):
    """Generate or display a specific chapter using the editor agent."""
    return _handle_chapter_stream(chapter_number, "editor")


@app.route("/inline_llm_continue_stream", methods=["POST"])
def inline_llm_continue_stream():
    """Get a streaming response from the LLM based on the provided context."""
    data = request.json
    context = data.get("context", "")

    if not context:
        return Response(
            json.dumps({"error": "No context provided"}),
            status=400,
            mimetype="application/json",
        )

    book_agents = BookAgents(agent_config)
    book_agents.create_agents("", 0)

    user_prompt_str = _fmt_user(
        prompts.INLINE_CONTINUE_PROMPT,
        context=context,
        user_input="",
        action_beats="",
    )

    stream = book_agents.generate_content_stream("inline_continuer", user_prompt_str)

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/inline_llm_revise_stream", methods=["POST"])
def inline_llm_revise_stream():
    """Get a streaming response from the LLM for revision."""
    data = request.json
    context = data.get("context", "")
    user_prompt = data.get("user_prompt", "")

    if not context:
        return Response(
            json.dumps({"error": "No context provided"}),
            status=400,
            mimetype="application/json",
        )

    book_agents = BookAgents(agent_config)
    book_agents.create_agents("", 0)

    user_prompt_str = _fmt_user(
        prompts.INLINE_REVISE_PROMPT,
        context=context,
        user_input=user_prompt,
        action_beats="",
    )

    stream = book_agents.generate_content_stream("inline_reviser", user_prompt_str)

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/save_chapter/<int:chapter_number>", methods=["POST"])
def save_chapter(chapter_number):
    """Save edited chapter content"""
    chapter_content = request.form.get("chapter_content")
    chapter_content = chapter_content.strip()
    chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{chapter_number}{TEXT_EXTENSION}")
    with open(chapter_path, "w", encoding='utf-8') as f:
        f.write(chapter_content)
    return jsonify({"success": True})


@app.route("/save_chapter_editor/<int:chapter_number>", methods=["POST"])
def save_chapter_editor(chapter_number):
    """Save edited chapter content (editor version)"""
    chapter_content = request.form.get("chapter_content")
    chapter_content = chapter_content.strip()
    chapter_path = os.path.join(
        CHAPTERS_DIR, f"chapter_{chapter_number}_editor{TEXT_EXTENSION}"
    )
    with open(chapter_path, "w", encoding='utf-8') as f:
        f.write(chapter_content)
    return jsonify({"success": True})


@app.route("/save_master_prompt", methods=["POST"])
def save_master_prompt():
    """Save the master prompt to a file."""
    master_prompt = request.form.get("master_prompt", "")
    with open(MASTER_PROMPT_FILE, "w", encoding='utf-8') as f:
        f.write(master_prompt)
    return jsonify({"success": True})


@app.route("/save_chapter_style/<int:chapter_number>", methods=["POST"])
def save_chapter_style(chapter_number):
    """Save chapter-specific style settings (point_of_view and tense)."""
    data = request.json
    point_of_view = data.get("point_of_view")
    tense = data.get("tense")
    min_words = data.get("min_words")

    settings = get_settings()
    if "chapters" not in settings:
        settings["chapters"] = {}
    if str(chapter_number) not in settings["chapters"]:
        settings["chapters"][str(chapter_number)] = {}

    if point_of_view is not None:
        settings["chapters"][str(chapter_number)]["point_of_view"] = point_of_view
    if tense is not None:
        settings["chapters"][str(chapter_number)]["tense"] = tense
    if min_words is not None:
        settings["chapters"][str(chapter_number)]["min_words"] = min_words

    save_settings(settings)
    return jsonify({"success": True})


@app.route("/save_setting", methods=["POST"])
def save_setting():
    """Save a specific setting value."""
    data = request.json
    key = data.get("key")
    value = data.get("value")

    if not key or value is None:
        return jsonify({"error": "Key or value missing"}), 400

    settings = get_settings()
    settings[key] = value
    save_settings(settings)
    return jsonify({"success": True})


@app.route("/scene/<int:chapter_number>", methods=["GET", "POST"])
def scene(chapter_number):
    """Generate a scene for a specific chapter"""
    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )

    if not chapter_data:
        chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{chapter_number}{TEXT_EXTENSION}")
        if os.path.exists(chapter_path):
            chapter_data = {
                "chapter_number": chapter_number,
                "title": f"Chapter {chapter_number}",
                "prompt": "Chapter content from file",
            }
        else:
            chapter_data = {
                "chapter_number": chapter_number,
                "title": f"Chapter {chapter_number}",
                "prompt": "No chapter outline available",
            }

    if request.method == "POST":
        world_theme = get_world_theme()
        characters = get_characters()
        previous_context = get_previous_chapter_context(chapter_number)

        book_agents = BookAgents(agent_config, chapters)
        book_agents.create_agents(world_theme, len(chapters) if chapters else 1)

        user_prompt_str = _fmt_user(
            prompts.SCENE_GENERATION_PROMPT,
            chapter_number=chapter_number,
            chapter_title=chapter_data.get("title", f"Chapter {chapter_number}"),
            chapter_outline=chapter_data.get("prompt", ""),
            world_theme=world_theme,
            relevant_characters=characters,
            previous_context=previous_context,
        )

        scene_content = book_agents.generate_content("writer", user_prompt_str)

        scene_dir = os.path.join(CHAPTERS_DIR, f"chapter_{chapter_number}_scenes")
        os.makedirs(scene_dir, exist_ok=True)

        scene_count = len([f for f in os.listdir(scene_dir) if f.endswith(TEXT_EXTENSION)])
        scene_path = os.path.join(scene_dir, f"scene_{scene_count + 1}{TEXT_EXTENSION}")

        with open(scene_path, "w", encoding='utf-8') as f:
            f.write(scene_content)

        return jsonify({"scene_content": scene_content})

    scenes = []
    scene_dir = os.path.join(CHAPTERS_DIR, f"chapter_{chapter_number}_scenes")

    if os.path.exists(scene_dir):
        scene_files = [f for f in os.listdir(scene_dir) if f.endswith(TEXT_EXTENSION)]
        scene_files.sort(key=lambda f: int(f.split("_")[1].split(".")[0]))

        for scene_file in scene_files:
            scene_path = os.path.join(scene_dir, scene_file)
            scene_number = int(scene_file.split("_")[1].split(".")[0])

            with open(scene_path, "r", encoding="utf-8") as f:
                content = f.read()
                lines = content.split("\n")
                title = lines[0][:30] + "..." if (lines and len(lines[0]) > 30) else (lines[0] if lines else f"Scene {scene_number}")
                scenes.append({"number": scene_number, "title": title, "content": content})

    return render_template("scene.html", chapter=chapter_data, scenes=scenes)


@app.route("/save_action_beats/<int:chapter_number>", methods=["POST"])
def save_action_beats(chapter_number):
    """Save edited action beats content"""
    action_beats_content = request.form.get("action_beats_content")
    action_beats_content = action_beats_content.strip()
    action_beats_path = os.path.join(
        CHAPTERS_DIR, f"chapter_{chapter_number}_action_beats{TEXT_EXTENSION}"
    )
    with open(action_beats_path, "w", encoding='utf-8') as f:
        f.write(action_beats_content)
    return jsonify({"success": True})


@app.route("/action_beats_chat/<int:chapter_number>", methods=["GET"])
def action_beats_chat(chapter_number):
    """Display action beats chat interface"""
    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )
    if not chapter_data:
        return render_template("error.html", message=f"Chapter {chapter_number} not found")
    action_beats_content = get_action_beats(chapter_number)
    return render_template(
        "action_beats_chat.html",
        chapter=chapter_data,
        action_beats_content=action_beats_content,
        chapters=chapters,
    )


@app.route("/action_beats_chat_stream/<int:chapter_number>", methods=["POST"])
def action_beats_chat_stream(chapter_number):
    """Handle ongoing chat for action beats creation with streaming response"""
    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )

    if not chapter_data:
        return Response(
            json.dumps({"error": f"Chapter {chapter_number} not found"}),
            status=404,
            mimetype="application/json",
        )

    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])

    world_theme = get_world_theme()
    characters = get_characters()

    book_agents = BookAgents(agent_config, chapters)
    book_agents.create_agents(world_theme, len(chapters) if chapters else 1)

    stream = book_agents.generate_chat_response_action_beats_stream(
        chat_history,
        chapter_data.get("prompt", ""),
        world_theme,
        characters,
        user_message,
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/finalize_action_beats_stream/<int:chapter_number>", methods=["POST"])
def finalize_action_beats_stream(chapter_number):
    """Finalize the action beats based on chat history with streaming response"""
    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )

    if not chapter_data:
        return Response(
            json.dumps({"error": f"Chapter {chapter_number} not found"}),
            status=404,
            mimetype="application/json",
        )

    data = request.json
    chat_history = data.get("chat_history", [])
    num_beats = data.get("num_beats", 12)

    world_theme = get_world_theme()
    characters = get_characters()

    book_agents = BookAgents(agent_config, chapters)
    book_agents.create_agents(world_theme, len(chapters) if chapters else 1)

    stream = book_agents.generate_final_action_beats_stream(
        chat_history,
        chapter_data.get("prompt", ""),
        world_theme,
        characters,
        num_beats,
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        collected_content = []
        for chunk in stream:
            content = chunk.content
            if content:
                collected_content.append(content)
                yield f"data: {json.dumps({'content': content})}\n\n"

        complete_content = "".join(collected_content)
        action_beats_path = os.path.join(
            CHAPTERS_DIR, f"chapter_{chapter_number}_action_beats{TEXT_EXTENSION}"
        )
        with open(action_beats_path, "w", encoding='utf-8') as f:
            f.write(complete_content)

        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/characters_chat", methods=["POST"])
def characters_chat():
    """Handle ongoing chat for character creation"""
    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])
    world_theme = get_world_theme()

    if not world_theme:
        return jsonify(
            {"error": "World theme not found. Please complete world building first."}
        )

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(world_theme, 0)

    ai_response = book_agents.generate_chat_response_characters(
        chat_history, world_theme, user_message
    )
    ai_response = ai_response.strip()
    return jsonify({"message": ai_response})


@app.route("/characters_chat_stream", methods=["POST"])
def characters_chat_stream():
    """Handle ongoing chat for character creation with streaming response"""
    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])
    world_theme = get_world_theme()

    if not world_theme:
        return jsonify(
            {"error": "World theme not found. Please complete world building first."}
        )

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(world_theme, 0)

    stream = book_agents.generate_chat_response_characters_stream(
        chat_history, world_theme, user_message
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/finalize_characters_stream", methods=["POST"])
def finalize_characters_stream():
    """Finalize the characters based on chat history with streaming response"""
    data = request.json
    chat_history = data.get("chat_history", [])
    num_characters = data.get("num_characters", 3)
    world_theme = get_world_theme()

    if not world_theme:
        return jsonify(
            {"error": "World theme not found. Please complete world building first."}
        )

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(world_theme, 0)

    stream = book_agents.generate_final_characters_stream(
        chat_history, world_theme, num_characters
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        collected_content = []
        for chunk in stream:
            content = chunk.content
            if content:
                collected_content.append(content)
                yield f"data: {json.dumps({'content': content})}\n\n"

        complete_content = "".join(collected_content)
        characters_content = complete_content.strip()

        with open(CHARACTERS_FILE, "w", encoding='utf-8') as f:
            f.write(characters_content)

        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/outline_chat", methods=["POST"])
def outline_chat():
    """Handle ongoing chat for outline creation"""
    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])
    num_chapters = data.get("num_chapters", 10)

    world_theme = get_world_theme()
    characters = get_characters()
    synopsis = get_synopsis()

    if not world_theme or not characters or not synopsis:
        return jsonify({
            "error": "World theme, characters, or synopsis not found. Please complete previous steps first."
        })

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(world_theme, num_chapters)

    ai_response = book_agents.generate_chat_response_outline(
        chat_history, world_theme, characters, synopsis, user_message
    )
    ai_response = ai_response.strip()
    return jsonify({"message": ai_response})


@app.route("/outline_chat_stream", methods=["POST"])
def outline_chat_stream():
    """Handle ongoing chat for outline creation with streaming response"""
    data = request.json
    user_message = data.get("message", "")
    chat_history = data.get("chat_history", [])
    num_chapters = data.get("num_chapters", 10)

    world_theme = get_world_theme()
    characters = get_characters()
    synopsis = get_synopsis()

    if not world_theme or not characters or not synopsis:
        return jsonify({
            "error": "World theme, characters, or synopsis not found. Please complete previous steps first."
        })

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(world_theme, num_chapters)

    stream = book_agents.generate_chat_response_outline_stream(
        chat_history, world_theme, characters, synopsis, user_message
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/finalize_outline_stream", methods=["POST"])
def finalize_outline_stream():
    """Finalize the outline based on chat history with streaming response"""
    data = request.json
    chat_history = data.get("chat_history", [])
    num_chapters = data.get("num_chapters", 10)

    world_theme = get_world_theme()
    characters = get_characters()
    synopsis = get_synopsis()

    if not world_theme or not characters or not synopsis:
        return jsonify({
            "error": "World theme, characters, or synopsis not found. Please complete previous steps first."
        })

    book_agents = BookAgents(agent_config)
    book_agents.create_agents(world_theme, num_chapters)

    stream = book_agents.generate_final_outline_stream(
        chat_history, world_theme, characters, synopsis, num_chapters
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        collected_content = []
        for chunk in stream:
            content = chunk.content
            if content:
                collected_content.append(content)
                yield f"data: {json.dumps({'content': content})}\n\n"

        complete_content = "".join(collected_content)
        outline_content = complete_content.strip()

        with open(OUTLINE_FILE, "w", encoding='utf-8') as f:
            f.write(outline_content)

        new_chapters = parse_outline_to_chapters(outline_content, num_chapters)
        with open(OUTLINE_JSON_FILE, "w", encoding="utf-8") as f:
            json.dump(new_chapters, f, indent=2)

        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


def parse_outline_to_chapters(outline_content, num_chapters):
    """Helper function to parse outline content into structured chapter format"""
    outline_content = outline_content.replace("\r\n", "\n")
    outline_content = re.sub(r"\n{2,}", "\n\n", outline_content)
    outline_content = outline_content.strip()

    chapters = []
    try:
        start_idx = outline_content.find("OUTLINE:")
        end_idx = outline_content.find("END OF OUTLINE")
        if start_idx != -1 and end_idx != -1:
            outline_text = outline_content[start_idx + len("OUTLINE:"): end_idx].strip()
        else:
            outline_text = outline_content

        outline_text = re.sub(r"###\s+\**Act\s+\w+:.+\n", "", outline_text)
        outline_text = re.sub(r"###\s+\**Epilogue\**.+\n", "", outline_text)

        chapter_matches = re.finditer(
            r"Chapter\s+(\d+):\s+\*?\*?([^\n*]+)\*?\*?", outline_text
        )
        seen_chapters = set()

        for match in chapter_matches:
            chapter_num = int(match.group(1))
            chapter_title = match.group(2).strip()

            if chapter_num in seen_chapters:
                continue
            seen_chapters.add(chapter_num)

            start_pos = match.start()
            next_chapter_match = re.search(
                r"Chapter\s+(\d+):", outline_text[start_pos + 1:]
            )

            if next_chapter_match:
                end_pos = start_pos + 1 + next_chapter_match.start()
                chapter_content = outline_text[start_pos:end_pos].strip()
            else:
                chapter_content = outline_text[start_pos:].strip()

            content_lines = chapter_content.split("\n")[1:]
            content_lines = [line.strip() for line in content_lines if line.strip()]
            chapter_description = "\n".join(content_lines).rstrip() if content_lines else ""

            chapter_description = re.sub(r"^\n+|\n+$", "", chapter_description)
            chapter_description = re.sub(r"\*+$", "", chapter_description)

            chapters.append({
                "chapter_number": chapter_num,
                "title": chapter_title,
                "prompt": chapter_description,
            })

        chapters.sort(key=lambda x: x["chapter_number"])

        if not chapters:
            for i in range(1, num_chapters + 1):
                chapters.append({
                    "chapter_number": i,
                    "title": f"Chapter {i}",
                    "prompt": f"Content for chapter {i}",
                })

    except Exception as e:
        print(f"Error parsing outline: {e}")
        for i in range(1, num_chapters + 1):
            chapters.append({
                "chapter_number": i,
                "title": f"Chapter {i}",
                "prompt": f"Content for chapter {i}",
            })

    with open(CHAPTERS_JSON_FILE, "w", encoding="utf-8") as f:
        json.dump(chapters, f, indent=2)

    return chapters


# API routes
@app.route("/api/chapters", methods=["GET"])
def api_chapters():
    """API endpoint to get all chapters."""
    all_chapters = get_chapters()
    return jsonify(all_chapters)


@app.route("/api/paginated_chapters", methods=["GET"])
def api_paginated_chapters():
    """API endpoint to get paginated chapters."""
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 10, type=int)
    data = get_paginated_chapters(page, per_page)
    return jsonify({
        "chapters": data["chapters"],
        "total_pages": data["total_pages"],
        "current_page": data["current_page"],
        "total_chapters": data["total_chapters"],
    })


@app.route("/api/models", methods=["POST"])
def api_models():
    """Fetch available models from the configured API endpoint."""
    data = request.json or {}
    base_url = data.get("base_url") or agent_config["config_list"][0]["base_url"]
    api_key = data.get("api_key") or agent_config["config_list"][0]["api_key"]
    try:
        from openai import OpenAI
        client = OpenAI(base_url=base_url, api_key=api_key)
        models = client.models.list()
        model_ids = sorted([m.id for m in models])
        return jsonify({"success": True, "models": model_ids})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/chapter/<int:chapter_number>", methods=["GET"])
def api_chapter(chapter_number):
    """API endpoint to get a specific chapter by number."""
    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )
    if not chapter_data:
        return Response(
            json.dumps({"error": f"Chapter {chapter_number} not found"}),
            status=404,
            mimetype="application/json",
        )
    return jsonify(chapter_data)


@app.route("/continuity_check_stream", methods=["POST"])
def continuity_check_stream():
    """Run a comprehensive continuity check across all manuscript content (streaming)."""
    characters = get_characters()
    world_theme = get_world_theme()
    synopsis = get_synopsis()

    # Collect all written chapter content
    chapters = get_chapters()
    chapters_parts = []
    for ch in chapters:
        chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{ch['chapter_number']}{TEXT_EXTENSION}")
        if os.path.exists(chapter_path):
            with open(chapter_path, "r", encoding="utf-8") as f:
                content = f.read().strip()
                if content:
                    chapters_parts.append(f"--- Chapter {ch['chapter_number']}: {ch['title']} ---\n{content}")

    chapters_content = "\n\n".join(chapters_parts) if chapters_parts else ""

    book_agents = BookAgents(agent_config)
    book_agents.create_agents("", 0)

    stream = book_agents.run_continuity_check_stream(
        characters, world_theme, synopsis, chapters_content
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/continuity_fix_stream", methods=["POST"])
def continuity_fix_stream():
    """Run the continuity fixer AI to resolve issues found in the last continuity report."""
    data = request.json
    continuity_report = data.get("continuity_report", "")

    characters = get_characters()
    world_theme = get_world_theme()
    synopsis = get_synopsis()

    chapters = get_chapters()
    chapters_parts = []
    for ch in chapters:
        chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{ch['chapter_number']}{TEXT_EXTENSION}")
        if os.path.exists(chapter_path):
            with open(chapter_path, "r", encoding="utf-8") as f:
                content = f.read().strip()
                if content:
                    chapters_parts.append(f"--- Chapter {ch['chapter_number']}: {ch['title']} ---\n{content}")

    chapters_content = "\n\n".join(chapters_parts) if chapters_parts else ""

    book_agents = BookAgents(agent_config)
    book_agents.create_agents("", 0)

    stream = book_agents.fix_continuity_issues_stream(
        continuity_report, characters, world_theme, synopsis, chapters_content
    )

    def generate():
        yield 'data: {"content": ""}\n\n'
        for chunk in stream:
            content = chunk.content
            if content:
                yield f"data: {json.dumps({'content': content})}\n\n"
        yield f"data: {json.dumps({'content': '[DONE]'})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/continuity_apply_fixes", methods=["POST"])
def continuity_apply_fixes():
    """Parse the fix report and apply the corrected text back into the chapter files."""
    data = request.json
    fix_report = data.get("fix_report", "")

    if not fix_report:
        return jsonify({"success": False, "error": "No fix report provided."}), 400

    # Load chapters index
    chapters = get_chapters()
    chapter_map = {}
    for ch in chapters:
        cn = ch["chapter_number"]
        chapter_map[cn] = ch
        chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{cn}{TEXT_EXTENSION}")
        if os.path.exists(chapter_path):
            with open(chapter_path, "r", encoding="utf-8") as f:
                ch["_content"] = f.read()

    # Parse the fix report for FIX blocks — line-by-line state-machine approach
    import re
    import textwrap

    applied_count = 0
    errors = []
    fixes_log = []

    # Split on any line containing "FIX #" with flexible formatting
    blocks = re.split(r'\n(?=---?\s*FIX\s*#?\d+\s*-?-?)', fix_report)
    # Also try splitting by "--- FIX" patterns
    if len(blocks) <= 1:
        blocks = re.split(r'---?\s*FIX\s*#?\d+', fix_report)

    if len(blocks) <= 1:
        return jsonify({"success": False, "error": "No FIX blocks found in the report. The AI may not have produced properly formatted output."}), 400

    for i, block in enumerate(blocks[1:], 1):
        lines = block.strip().split('\n')

        # Extract chapter number — look for "Chapter: N" or "Chapter N:" patterns
        chapter_num = None
        old_text_start = None
        new_text_start = None

        for j, line in enumerate(lines):
            stripped = line.strip()
            # Check for chapter number
            if chapter_num is None:
                ch_match = re.search(r'Chapter\s*:?\s*(\d+)', stripped, re.IGNORECASE)
                if ch_match:
                    chapter_num = int(ch_match.group(1))
            # Find Old Text / Corrected Text boundaries
            if re.match(r'Old\s+Text\s*:?', stripped, re.IGNORECASE):
                old_text_start = j + 1
            if re.match(r'Corrected\s+Text\s*:?', stripped, re.IGNORECASE):
                new_text_start = j + 1

        if chapter_num is None:
            errors.append(f"Fix #{i}: Could not determine chapter from block: {block[:100]}...")
            continue
        if old_text_start is None:
            errors.append(f"Fix #{i}: Could not find Old Text marker")
            continue
        if new_text_start is None or new_text_start <= old_text_start:
            errors.append(f"Fix #{i}: Could not find Corrected Text marker after Old Text")
            continue

        # Extract old text (from Old Text marker to Corrected Text marker)
        old_text = '\n'.join(lines[old_text_start:new_text_start - 1]).strip()
        # Extract new text (from Corrected Text marker to end or next marker)
        # Stop at "--- FIX" or "FIXING COMPLETE" or end
        new_text_lines = []
        for line in lines[new_text_start:]:
            if re.match(r'---?\s*FIX|FIXING\s+COMPLETE|UNABLE\s+TO\s+FIX', line, re.IGNORECASE):
                break
            new_text_lines.append(line)
        new_text = '\n'.join(new_text_lines).strip()

        if not old_text:
            errors.append(f"Fix #{i}: Old Text is empty")
            continue
        if not new_text:
            errors.append(f"Fix #{i}: Corrected Text is empty")
            continue

        if chapter_num not in chapter_map:
            errors.append(f"Chapter {chapter_num} not found in the manuscript")
            continue

        ch = chapter_map[chapter_num]
        content = ch.get("_content", "")
        if not content:
            errors.append(f"Chapter {chapter_num} has no content")
            continue

        # Try exact match, then normalized whitespace match
        applied = False
        for text_variant, new_variant in [
            (old_text, new_text),
            (textwrap.dedent(old_text).strip(), textwrap.dedent(new_text).strip()),
            (' '.join(old_text.split()), ' '.join(new_text.split())),
        ]:
            if text_variant and text_variant in content:
                content = content.replace(text_variant, new_variant, 1)
                applied_count += 1
                fixes_log.append(f"Chapter {chapter_num}: replaced old text with corrected text")
                applied = True
                break

        if not applied:
            errors.append(f"Chapter {chapter_num}: Old Text not found in chapter content (may have already been changed or quoted differently)")

        # Save the chapter
        chapter_path = os.path.join(CHAPTERS_DIR, f"chapter_{chapter_num}{TEXT_EXTENSION}")
        with open(chapter_path, "w", encoding="utf-8") as f:
            f.write(content)

    return jsonify({
        "success": True,
        "applied_count": applied_count,
        "errors": errors,
        "fixes_log": fixes_log,
    })


@app.route("/research_brief/<int:chapter_number>", methods=["POST"])
def research_brief(chapter_number):
    """Generate a research brief for a chapter using SearXNG web search."""
    from config import _load_config_file

    raw_cfg = _load_config_file()
    research_enabled = raw_cfg.get("research_agent_enabled", False)
    searxng_host = raw_cfg.get("searxng_host", "").strip()

    if not research_enabled:
        return jsonify({
            "error": "Research agent is disabled. Enable it in Configuration."
        }), 400

    if not searxng_host:
        return jsonify({
            "error": "SearXNG host not configured. Set it in Configuration."
        }), 400

    chapters = get_chapters()
    chapter_data = next(
        (ch for ch in chapters if ch["chapter_number"] == chapter_number), None
    )

    if not chapter_data:
        return jsonify({"error": f"Chapter {chapter_number} not found"}), 404

    world_theme = get_world_theme()
    characters = get_characters()

    book_agents = BookAgents(agent_config, chapters)
    book_agents.create_agents(world_theme, len(chapters) if chapters else 1)

    try:
        search_results = book_agents.fetch_research_results(
            searxng_host,
            chapter_data["title"],
            chapter_data.get("prompt", ""),
            world_theme,
            characters,
        )
    except Exception as e:
        return jsonify({
            "error": f"Failed to fetch search results from SearXNG: {str(e)}"
        }), 502

    try:
        brief = book_agents.generate_research_brief(
            chapter_data["title"],
            chapter_data.get("prompt", ""),
            world_theme,
            characters,
            search_results,
        )
    except Exception as e:
        return jsonify({
            "error": f"Failed to generate research brief: {str(e)}"
        }), 500

    # Persist per-chapter brief
    brief_path = os.path.join(
        CHAPTERS_DIR,
        f"chapter_{chapter_number}_research_brief{TEXT_EXTENSION}",
    )
    with open(brief_path, "w", encoding="utf-8") as f:
        f.write(brief)

    # Return brief + search result count
    result_count = len(search_results) if search_results else 0
    return jsonify({
        "success": True,
        "research_brief": brief,
        "search_results_count": result_count,
        "queries_used": getattr(book_agents, '_last_queries', []),
    })


if __name__ == "__main__":
    check_openai_connection(agent_config)
    if str(agent_config.get("debug", False)).lower() in ("true", "1", "t"):
        print("=" * 50)
        print("🚀 CAUTION: DEBUG mode is enabled.")
        print(f"📝 Prompts, requests, and responses will be saved to the '{PROMPT_DEBUGGING_DIR}' directory.")
        print("=" * 50)
    app.run(debug=True)